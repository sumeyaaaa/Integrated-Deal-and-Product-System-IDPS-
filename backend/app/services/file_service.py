"""
File Upload and Analysis Service
=================================

Handles file uploads to Supabase storage and text extraction from various file types.
"""

import io
import mimetypes
from pathlib import Path
from typing import Optional, Tuple
from uuid import uuid4

from supabase import Client

from app.database.connection import get_supabase_client
from app.config import settings

# File processing imports
import PyPDF2
import pdfplumber
from docx import Document
import openpyxl
import pandas as pd


def ensure_bucket_exists(bucket_name: str = "attached_FILES", is_public: bool = True) -> None:
    """
    Ensure the storage bucket exists. Create it if it doesn't.
    
    Args:
        bucket_name: Name of the bucket
        is_public: Whether the bucket should be public (default: True)
    """
    from app.database.connection import get_supabase_service_client
    storage_client = get_supabase_service_client()
    
    try:
        # Try to list buckets to check if it exists
        buckets = storage_client.storage.list_buckets()
        bucket_exists = any(b.name == bucket_name for b in buckets)
        
        if not bucket_exists:
            # Create the bucket
            storage_client.storage.create_bucket(
                bucket_name,
                options={"public": is_public}
            )
    except Exception as e:
        # If bucket already exists or other error, that's okay
        # We'll try to upload anyway
        pass


def upload_file_to_supabase(
    file_content: bytes,
    filename: str,
    bucket_name: str = "attached_FILES",
) -> Tuple[str, str]:
    """
    Upload a file to Supabase storage bucket.
    
    Args:
        file_content: File content as bytes
        filename: Original filename
        bucket_name: Supabase storage bucket name (default: "attached_FILES")
    
    Returns:
        Tuple of (file_url, file_type)
    """
    # Ensure bucket exists before uploading
    ensure_bucket_exists(bucket_name, is_public=True)
    
    # Generate unique filename to avoid conflicts
    file_ext = Path(filename).suffix
    unique_filename = f"{uuid4()}{file_ext}"
    
    # Determine content type
    content_type, _ = mimetypes.guess_type(filename)
    if not content_type:
        content_type = "application/octet-stream"
    
    # Upload to Supabase storage
    try:
        # Use service client for storage operations (needs admin access)
        from app.database.connection import get_supabase_service_client
        storage_client = get_supabase_service_client()
        
        response = storage_client.storage.from_(bucket_name).upload(
            unique_filename,
            file_content,
            file_options={"content-type": content_type, "upsert": "false"}
        )
        
        # Get public URL - construct it from Supabase URL
        from app.config import settings
        supabase_url = settings.SUPABASE_URL.rstrip("/")
        file_url = f"{supabase_url}/storage/v1/object/public/{bucket_name}/{unique_filename}"
        
        return file_url, content_type
    except Exception as e:
        raise RuntimeError(f"Failed to upload file to Supabase: {str(e)}")


def extract_text_from_file(file_content: bytes, filename: str, content_type: str) -> str:
    """
    Extract text content from various file types.
    
    Supports:
    - PDF (PyPDF2 and pdfplumber)
    - Word documents (.docx)
    - Excel files (.xlsx, .xls)
    - Text files (.txt)
    
    Args:
        file_content: File content as bytes
        filename: Original filename
        content_type: MIME type of the file
    
    Returns:
        Extracted text content
    """
    file_ext = Path(filename).suffix.lower()
    text_parts = []
    
    try:
        if file_ext == ".pdf" or content_type == "application/pdf":
            # Try pdfplumber first (better for complex PDFs)
            try:
                pdf_file = io.BytesIO(file_content)
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
            except Exception:
                # Fallback to PyPDF2
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        
        elif file_ext == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
        
        elif file_ext in [".xlsx", ".xls"] or "spreadsheet" in content_type:
            excel_file = io.BytesIO(file_content)
            # Try openpyxl for .xlsx
            if file_ext == ".xlsx":
                workbook = openpyxl.load_workbook(excel_file)
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text_parts.append(f"\n--- Sheet: {sheet_name} ---\n")
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        if row_text.strip():
                            text_parts.append(row_text)
            else:
                # Use pandas for .xls
                df = pd.read_excel(excel_file, sheet_name=None)
                for sheet_name, sheet_df in df.items():
                    text_parts.append(f"\n--- Sheet: {sheet_name} ---\n")
                    text_parts.append(sheet_df.to_string())
        
        elif file_ext == ".txt" or content_type == "text/plain":
            text_parts.append(file_content.decode("utf-8", errors="ignore"))
        
        else:
            # For other file types, return a note
            text_parts.append(f"[File type {file_ext} not directly readable. File uploaded: {filename}]")
    
    except Exception as e:
        text_parts.append(f"[Error extracting text from {filename}: {str(e)}]")
    
    return "\n\n".join(text_parts) if text_parts else f"[No text extracted from {filename}]"

